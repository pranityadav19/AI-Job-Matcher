"""
Resume Text Extractor
Extracts text from various resume formats (PDF, DOCX, TXT)

IMPROVEMENTS IN THIS VERSION:
- Better error handling with fallback methods
- More informative error messages
- Handles corrupted files gracefully
"""

import os
from typing import Optional
import PyPDF2
from docx import Document


class ResumeExtractor:
    """Extract text from resume files in various formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF has pages
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file appears to be empty")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {str(e)}")
                        continue
                
                if not text.strip():
                    raise ValueError("No text could be extracted from PDF. It may be an image-based PDF.")
                    
        except Exception as e:
            raise ValueError(f"Error reading PDF file: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            
            if not text.strip():
                raise ValueError("DOCX file appears to be empty")
                
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_from_txt(file_path: str) -> str:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                # Try latin-1 encoding
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except UnicodeDecodeError:
                # Last resort: ignore errors
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                print("Warning: Some characters could not be decoded and were skipped")
        except Exception as e:
            raise ValueError(f"Error reading TXT file: {str(e)}")
        
        if not text.strip():
            raise ValueError("TXT file appears to be empty")
        
        return text.strip()
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from resume file (auto-detects format)
        
        IMPROVED: Better error handling and fallback methods
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError(f"File is empty: {file_path}")
        
        if file_size > 50 * 1024 * 1024:  # 50 MB
            print(f"Warning: Large file detected ({file_size / 1024 / 1024:.1f} MB). This may take a while...")
        
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Extract based on file type
        try:
            if ext == '.pdf':
                return ResumeExtractor.extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return ResumeExtractor.extract_from_docx(file_path)
            elif ext == '.txt':
                return ResumeExtractor.extract_from_txt(file_path)
            else:
                # IMPROVED: Try fallback for unknown extensions
                print(f"Warning: Unsupported extension '{ext}'. Trying to read as plain text...")
                try:
                    return ResumeExtractor.extract_from_txt(file_path)
                except:
                    raise ValueError(f"Unsupported file format: {ext}. Supported formats: PDF, DOCX, TXT")
                    
        except Exception as e:
            # IMPROVED: Provide helpful error message
            error_msg = f"Could not extract text from {file_path}:\n{str(e)}\n\n"
            error_msg += "Troubleshooting tips:\n"
            error_msg += "- Ensure the file is not corrupted\n"
            error_msg += "- For PDFs: Make sure it's not an image-based PDF (use OCR if needed)\n"
            error_msg += "- For DOCX: Try opening and re-saving in Microsoft Word\n"
            error_msg += "- Try converting your file to TXT format\n"
            raise ValueError(error_msg)


def main():
    """Example usage"""
    import sys
    
    if len(sys.argv) < 2:
        print("="*60)
        print("RESUME TEXT EXTRACTOR")
        print("="*60)
        print("\nUsage: python ResumeExtractor.py <resume_file>")
        print("\nSupported formats: PDF, DOCX, DOC, TXT")
        print("\nExamples:")
        print("  python ResumeExtractor.py my_resume.pdf")
        print("  python ResumeExtractor.py resume.docx")
        print("="*60)
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    print("="*60)
    print("RESUME TEXT EXTRACTOR")
    print("="*60)
    print(f"\nFile: {file_path}")
    
    try:
        extractor = ResumeExtractor()
        print("Extracting text...")
        text = extractor.extract_text(file_path)
        
        print("\n" + "="*60)
        print("EXTRACTED RESUME TEXT")
        print("="*60 + "\n")
        print(text)
        print("\n" + "="*60)
        print("EXTRACTION STATISTICS")
        print("="*60)
        print(f"Total characters: {len(text):,}")
        print(f"Total words: {len(text.split()):,}")
        print(f"Total lines: {len(text.splitlines()):,}")
        print("="*60)
        print("\n[OK] Extraction successful!")
        
    except Exception as e:
        print("\n" + "="*60)
        print("ERROR")
        print("="*60)
        print(f"\n{str(e)}")
        print("\n" + "="*60)
        sys.exit(1)


if __name__ == "__main__":
    main()